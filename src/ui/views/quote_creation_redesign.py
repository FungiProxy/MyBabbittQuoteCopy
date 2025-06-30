"""
Redesigned Quote Creation Page

Clean, step-by-step quote creation interface with integrated product selection,
configuration, and customer information forms. Focuses on user-friendly workflow.
"""

import logging
from typing import Dict, List, Optional
import os

from PySide6.QtCore import Qt, Signal
from PySide6.QtWidgets import (
    QFrame,
    QGridLayout,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QLineEdit,
    QPushButton,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QVBoxLayout,
    QWidget,
    QMessageBox,
    QGroupBox,
    QFormLayout,
    QSizePolicy,
    QGraphicsDropShadowEffect,
    QAbstractItemView,
    QSpacerItem,
    QScrollArea,
    QStackedWidget,
    QFileDialog,
)
from PySide6.QtGui import QColor, QFont, QIcon, QPalette

from src.core.database import SessionLocal
from src.core.services.quote_service import QuoteService
from src.core.services.product_service import ProductService
from src.core.models.quote import Quote
from src.ui.dialogs.customer_selection_dialog import CustomerSelectionDialog
from src.ui.dialogs.customer_dialog import CustomerDialog
from src.ui.product_selection_dialog_modern import ModernProductSelectionDialog
from src.ui.theme.babbitt_theme import BabbittTheme
from src.ui.theme.theme_manager import ThemeManager
from src.ui.theme import COLORS, FONTS

logger = logging.getLogger(__name__)


class QuoteCreationPageRedesign(QWidget):
    """
    Redesigned quote creation page with clean workflow and modern interface.
    
    Features:
    - Step-by-step quote building
    - Integrated product selection and configuration
    - Real-time quote totals
    - Customer information management
    - Professional quote summary
    """
    
    quote_created = Signal(dict)  # Emitted when quote is successfully created
    quote_updated = Signal(dict)  # Emitted when quote is updated
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # Services
        self.db = SessionLocal()
        self.quote_service = QuoteService()
        self.product_service = ProductService(db=self.db)
        
        # Current quote state
        self.current_quote_id = None
        self.current_quote = {
            "items": [],
            "customer_info": {},
            "total_value": 0.0,
            "quote_number": None,
            "status": "Draft"
        }
        
        self._setup_ui()
        self._update_items_visibility()

    def __del__(self):
        """Clean up database connection."""
        if hasattr(self, 'db') and self.db:
            self.db.close()

    def _style_input_field(self, field, uniform_height=48):
        field.setStyleSheet(f"""
            QLineEdit {{
                border: 2px solid #E9ECEF;
                border-radius: 6px;
                padding: 12px 16px;
                font-size: 14px;
                background: white;
                height: {uniform_height}px;
            }}
            QLineEdit:focus {{
                border-color: #2C3E50;
                outline: none;
            }}
            QLineEdit:hover {{
                border-color: #34495E;
            }}
        """)
    
    def _style_text_area(self, field):
        field.setStyleSheet("""
            QTextEdit {
                border: 2px solid #E9ECEF;
                border-radius: 6px;
                padding: 12px 16px;
                font-size: 14px;
                background: white;
                min-height: 120px;
            }
            QTextEdit:focus {
                border-color: #2C3E50;
                outline: none;
            }
            QTextEdit:hover {
                border-color: #34495E;
            }
        """)

    def _setup_ui(self):
        """Set up the quote creation UI to match the desired modern look."""
        # Main layout is now a QHBoxLayout for the two-column structure
        main_layout = QHBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        # --- Left Column (2/3 width) ---
        left_column_layout = QVBoxLayout()
        left_column_layout.setSpacing(20)
        
        # Create and add the header for the left column
        header_widget = self._create_quote_header()
        left_column_layout.addWidget(header_widget)
        
        # Create and add the items panel
        items_panel = self._create_items_panel()
        left_column_layout.addWidget(items_panel)
        
        # --- Right Column (1/3 width) ---
        right_column_layout = QVBoxLayout()
        right_column_layout.setSpacing(20)
        
        # Create and add the customer panel
        customer_panel = self._create_customer_panel()
        right_column_layout.addWidget(customer_panel)
        
        # Create and add the actions panel
        actions_panel = self._create_actions_panel()
        right_column_layout.addWidget(actions_panel)
        
        # Spacer to push customer/actions panels up
        right_column_layout.addStretch(1)
        
        # Add the two columns to the main layout with the correct stretch factor
        main_layout.addLayout(left_column_layout, 2)
        main_layout.addLayout(right_column_layout, 1)

    def mousePressEvent(self, event):
        """Clear focus from input widgets when clicking on the background."""
        focused_widget = self.focusWidget()
        if focused_widget and isinstance(focused_widget, (QLineEdit, QTextEdit)):
            focused_widget.clearFocus()
        super().mousePressEvent(event)

    def _create_quote_header(self) -> QWidget:
        """Create the quote header with title and basic info."""
        header = QWidget()
        layout = QHBoxLayout(header)
        layout.setContentsMargins(0, 0, 0, 0)

        # Left-aligned content: Quote Number
        self.quote_number_label = QLabel("Quote not yet saved")
        self.quote_number_label.setProperty("labelType", "caption")
        self.quote_number_label.setStyleSheet(f"""
            font-size: {FONTS['sizes']['lg']}px;
            font-weight: {FONTS['weights']['medium']};
            color: {COLORS['text_primary']};
            padding: 0 {FONTS['sizes']['sm']}px 0 0;
        """)

        # Center-aligned content: Status and Price
        status_price_container = QWidget()
        status_layout = QVBoxLayout(status_price_container)
        status_layout.setContentsMargins(0,0,0,0)
        status_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        self.status_label = QLabel(self.current_quote['status'].upper())
        self.status_label.setProperty("status", self.current_quote['status'])
        self.status_label.setStyleSheet(f"""
            font-size: {FONTS['sizes']['xl']}px;
            font-weight: {FONTS['weights']['bold']};
            color: {COLORS['primary']};
            letter-spacing: 1px;
            padding-bottom: 4px;
        """)
        status_layout.addWidget(self.status_label)
        
        self.total_label = QLabel(f"${self.current_quote['total_value']:.2f}")
        self.total_label.setProperty("priceType", "total-prominent")
        self.total_label.setStyleSheet(f"""
            font-size: {FONTS['sizes']['3xl']}px;
            font-weight: {FONTS['weights']['bold']};
            color: {COLORS['success']};
            padding-top: 2px;
        """)
        status_layout.addWidget(self.total_label)

        layout.addWidget(self.quote_number_label, 0, Qt.AlignmentFlag.AlignLeft)
        layout.addWidget(status_price_container, 0, Qt.AlignmentFlag.AlignCenter)
        layout.addStretch()
        return header

    def _apply_shadow_effect(self, widget):
        """Apply a standard drop shadow effect to a widget."""
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(20)
        shadow.setColor(QColor(0, 0, 0, 80))
        shadow.setOffset(0, 4)
        widget.setGraphicsEffect(shadow)

    def _create_items_panel(self) -> QWidget:
        """Create the quote items panel as a styled card."""
        panel = QFrame()
        panel.setProperty("class", "content-section")
        self._apply_shadow_effect(panel)

        layout = QVBoxLayout(panel)
        layout.setSpacing(15)

        # Title
        title_label = QLabel("Quote Items")
        title_label.setProperty("class", "section-title")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        # Container for the centered button
        button_container = QWidget()
        button_layout = QHBoxLayout(button_container)
        button_layout.addStretch()
        add_product_btn = QPushButton("+ Add Product")
        add_product_btn.clicked.connect(self._add_product)
        button_layout.addWidget(add_product_btn)
        button_layout.addStretch()
        layout.addWidget(button_container)
        
        # Stack for items table and empty state
        self.items_stack = QStackedWidget()
        
        # Items table
        self.items_table = QTableWidget()
        self.items_table.setColumnCount(6)
        self.items_table.setHorizontalHeaderLabels([
            "Product", "Configuration", "Quantity", "Unit Price", "Total", "Actions"
        ])
        
        header = self.items_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)
        
        self.items_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.items_table.setAlternatingRowColors(True)
        
        size_policy = self.items_table.sizePolicy()
        size_policy.setVerticalPolicy(QSizePolicy.Policy.Expanding)
        self.items_table.setSizePolicy(size_policy)
        self.items_stack.addWidget(self.items_table)

        # Empty state message
        empty_state_widget = QWidget()
        empty_layout = QVBoxLayout(empty_state_widget)
        empty_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.empty_state_label = QLabel("No items added yet. Click 'Add Product' to get started.")
        self.empty_state_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.empty_state_label.setProperty("class", "placeholderCard")
        empty_layout.addWidget(self.empty_state_label)
        self.items_stack.addWidget(empty_state_widget)
        
        layout.addWidget(self.items_stack)
        
        return panel

    def _create_customer_panel(self) -> QWidget:
        """Create the customer information panel as a styled card."""
        panel = QFrame()
        panel.setProperty("class", "content-section")
        self._apply_shadow_effect(panel)

        layout = QFormLayout(panel)

        # Action buttons for customer
        action_layout = QHBoxLayout()
        select_customer_btn = QPushButton("Select Customer")
        select_customer_btn.clicked.connect(self._select_customer)
        new_customer_btn = QPushButton("New Customer")
        new_customer_btn.clicked.connect(self._add_new_customer)
        action_layout.addWidget(select_customer_btn)
        action_layout.addWidget(new_customer_btn)
        layout.addRow(action_layout)

        # Form fields
        self.company_name_edit = QLineEdit()
        self.company_name_edit.setPlaceholderText("Company Name")
        
        self.contact_person_edit = QLineEdit()
        self.contact_person_edit.setPlaceholderText("Contact Person")
        
        self.email_edit = QLineEdit()
        self.email_edit.setPlaceholderText("Email")
        
        self.phone_edit = QLineEdit()
        self.phone_edit.setPlaceholderText("Phone")
        
        self.notes_edit = QTextEdit()
        self.notes_edit.setPlaceholderText("Notes")
        
        layout.addRow("Company Name:", self.company_name_edit)
        layout.addRow("Contact Person:", self.contact_person_edit)
        layout.addRow("Email:", self.email_edit)
        layout.addRow("Phone:", self.phone_edit)
        layout.addRow("Notes:", self.notes_edit)
        
        return panel

    def _create_actions_panel(self) -> QWidget:
        """Create the quote actions panel with save, export, and send buttons."""
        panel = QFrame()
        panel.setProperty("class", "content-section")
        self._apply_shadow_effect(panel)

        layout = QVBoxLayout(panel)
        layout.setSpacing(15)

        title_label = QLabel("Actions")
        title_label.setProperty("class", "section-title")
        layout.addWidget(title_label)

        self.save_draft_btn = QPushButton("Save Draft")
        self.save_draft_btn.clicked.connect(self._save_draft)
        layout.addWidget(self.save_draft_btn)

        self.generate_word_btn = QPushButton("Export to Word")
        self.generate_word_btn.setProperty("class", "secondary")
        self.generate_word_btn.clicked.connect(self._export_to_word)
        layout.addWidget(self.generate_word_btn)

        # Add a separator and a clear button
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        layout.addWidget(separator)

        self.clear_quote_btn = QPushButton("Clear Quote")
        self.clear_quote_btn.setProperty("class", "danger")
        self.clear_quote_btn.clicked.connect(self.new_quote)
        layout.addWidget(self.clear_quote_btn)

        return panel

    def _select_customer(self):
        """Open a dialog to select an existing customer."""
        dialog = CustomerSelectionDialog(self)
        dialog.customer_selected.connect(self._handle_customer_selected)
        dialog.exec()

    def _add_new_customer(self):
        """Open a dialog to create a new customer."""
        dialog = CustomerDialog(self)
        dialog.customer_saved.connect(self._handle_customer_selected)
        dialog.exec()

    def _handle_customer_selected(self, customer_data):
        """Populate customer fields with data from selection."""
        self.current_quote["customer_info"] = customer_data
        
        self.company_name_edit.setText(customer_data.get("company", ""))
        self.contact_person_edit.setText(customer_data.get("name", ""))
        self.email_edit.setText(customer_data.get("email", ""))
        self.phone_edit.setText(customer_data.get("phone", ""))
        
        # You may need to fetch and set address fields if they exist
        # and are returned by the dialog.

    def _add_product(self):
        """Show the product selection dialog."""
        db = SessionLocal()
        try:
            # Use the modern dialog
            dialog = ModernProductSelectionDialog(product_service=self.product_service, parent=self)
            dialog.product_added.connect(self._on_product_configured)
            dialog.exec()
        except Exception as e:
            logger.error(f"Error opening product selection dialog: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Could not open product selection dialog: {e}")

    def _on_product_configured(self, config_data: Dict):
        """Handle completed product configuration from the new widget."""
        try:
            # The new widget returns a summary with all the data we need.
            quote_item = {
                "product_id": config_data.get("product_id"), # Make sure this is passed
                "product_family": config_data.get("product", "N/A"),
                "model_number": config_data.get("product", "N/A"), # Use the base product name
                "configuration": config_data.get("description", "Standard Configuration"),
                "quantity": config_data.get("quantity", 1),
                "unit_price": config_data.get("unit_price", 0),
                "total_price": config_data.get("total_price", 0),
                "config_data": config_data.get("configuration", {}), # Store the detailed config
                "options": config_data.get("options", []) # Make sure this is passed
            }
            
            self.current_quote["items"].append(quote_item)
            self._update_items_table()
            self._update_quote_totals()
            
            logger.info(f"Added product to quote: {quote_item['model_number']}")
            
        except Exception as e:
            logger.error(f"Error adding configured product: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Failed to add product to quote: {str(e)}")

    def _format_configuration(self, selected_options: Dict) -> str:
        """Format configuration options for display."""
        config_parts = []
        
        # Key configuration elements
        material = selected_options.get("Material")
        if material:
            config_parts.append(f"Material: {material}")
        
        voltage = selected_options.get("Voltage")
        if voltage:
            config_parts.append(f"Voltage: {voltage}")
        
        length = selected_options.get("Probe Length")
        if length:
            config_parts.append(f"Length: {length}\"")
        
        connection = selected_options.get("Connection")
        if connection and connection != "NPT":
            config_parts.append(f"Connection: {connection}")
        
        # Additional options
        if selected_options.get("Extended Probe"):
            config_parts.append("Extended Probe")
        
        if selected_options.get("Explosion Proof"):
            config_parts.append("Explosion Proof")
        
        if selected_options.get("NEMA 4X"):
            config_parts.append("NEMA 4X")
        
        return " | ".join(config_parts) if config_parts else "Standard Configuration"

    def _update_items_table(self):
        """Update the items table with current quote items."""
        # Safety check: ensure current_quote exists and has items
        if not hasattr(self, 'current_quote') or not self.current_quote:
            self.current_quote = {
                "items": [],
                "customer_info": {},
                "total_value": 0.0,
                "quote_number": None,
                "status": "Draft"
            }
        
        items = self.current_quote.get("items", [])
        self.items_table.setRowCount(len(items))
        
        for row, item in enumerate(items):
            # Product name
            product_item = QTableWidgetItem(item.get("product_family", "N/A"))
            product_item.setFlags(product_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.items_table.setItem(row, 0, product_item)
            
            # Configuration
            config_item = QTableWidgetItem(item.get("configuration", "Standard Configuration"))
            config_item.setFlags(config_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.items_table.setItem(row, 1, config_item)
            
            # Quantity (editable)
            quantity_item = QTableWidgetItem(str(item.get("quantity", 1)))
            quantity_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.items_table.setItem(row, 2, quantity_item)
            
            # Unit price
            unit_price_item = QTableWidgetItem(f"${item.get('unit_price', 0):,.2f}")
            unit_price_item.setFlags(unit_price_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            unit_price_item.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.items_table.setItem(row, 3, unit_price_item)
            
            # Total price
            total_price_item = QTableWidgetItem(f"${item.get('total_price', 0):,.2f}")
            total_price_item.setFlags(total_price_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            total_price_item.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.items_table.setItem(row, 4, total_price_item)
            
            # Actions
            actions_widget = self._create_actions_widget(row)
            self.items_table.setCellWidget(row, 5, actions_widget)
        
        # Connect quantity change signal
        self.items_table.itemChanged.connect(self._on_item_changed)
        
        self._update_items_visibility()

    def _create_actions_widget(self, row: int) -> QWidget:
        """Create action buttons for table row."""
        widget = QWidget()
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(5, 2, 5, 2)
        layout.setSpacing(5)
        
        # Edit button
        edit_btn = QPushButton("Edit")
        edit_btn.setProperty("class", "info small")
        
        # Remove button
        remove_btn = QPushButton("Remove")
        remove_btn.setProperty("class", "danger small")
        
        layout.addWidget(edit_btn)
        layout.addWidget(remove_btn)
        
        return widget

    def _update_items_visibility(self):
        """Update visibility of items table vs empty state."""
        has_items = len(self.current_quote["items"]) > 0
        self.items_stack.setCurrentIndex(0 if has_items else 1)
        
        # Update action buttons
        self.save_draft_btn.setEnabled(has_items)
        self.generate_word_btn.setEnabled(has_items and self._has_customer_info())
        self.clear_quote_btn.setEnabled(has_items)

    def _has_customer_info(self) -> bool:
        """Check if customer information is filled in."""
        return bool(
            self.company_name_edit.text().strip() and
            self.contact_person_edit.text().strip() and
            self.email_edit.text().strip()
        )

    def _on_item_changed(self, item: QTableWidgetItem):
        """Handle changes to table items (e.g., quantity)."""
        if item.column() == 2:  # Quantity column
            try:
                row = item.row()
                new_quantity = int(item.text())
                
                if new_quantity <= 0:
                    raise ValueError("Quantity must be positive")
                
                # Update quote item
                quote_item = self.current_quote["items"][row]
                quote_item["quantity"] = new_quantity
                quote_item["total_price"] = quote_item["unit_price"] * new_quantity
                
                # Update total price cell
                total_item = QTableWidgetItem(f"${quote_item['total_price']:,.2f}")
                total_item.setFlags(total_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                total_item.setTextAlignment(Qt.AlignmentFlag.AlignRight)
                self.items_table.setItem(row, 4, total_item)
                
                self._update_quote_totals()
                
            except ValueError:
                # Reset to previous value
                item.setText("1")
                QMessageBox.warning(self, "Invalid Quantity", "Please enter a valid positive number.")

    def _edit_item(self, row: int):
        """Edit an existing quote item."""
        product_to_edit = self.current_quote["items"][row]
        
        try:
            # Use the modern dialog for editing
            dialog = ModernProductSelectionDialog(
                product_service=self.product_service,
                product_to_edit=product_to_edit,
                parent=self
            )
            
            # Reconnect the signal to a handler that updates the item
            dialog.product_added.connect(lambda new_config: self._on_product_updated(row, new_config))
            
            dialog.exec()
        except Exception as e:
            logger.error(f"Error opening product configuration dialog for editing: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Could not open product configuration: {e}")

    def _on_product_updated(self, row: int, new_config: Dict):
        """
        Handle the updated configuration for an existing item.
        This replaces the old item at the given row with the new one.
        """
        # Replace the item in the quote data
        self.current_quote["items"][row] = new_config
        
        # Refresh the table to show updated item
        self._update_items_table()

    def _remove_item(self, row: int):
        """Remove an item from the quote."""
        if not (0 <= row < len(self.current_quote["items"])):
            return
        
        quote_item = self.current_quote["items"][row]
        
        reply = QMessageBox.question(
            self,
            "Remove Item",
            f"Are you sure you want to remove {quote_item['model_number']} from the quote?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.current_quote["items"].pop(row)
            self._update_items_table()
            self._update_quote_totals()

    def _update_quote_totals(self):
        """Update the quote totals display."""
        # Safety check: ensure current_quote exists
        if not hasattr(self, 'current_quote') or not self.current_quote:
            self.current_quote = {
                "items": [],
                "customer_info": {},
                "total_value": 0.0,
                "quote_number": None,
                "status": "Draft"
            }
        
        total = sum(item.get("total_price", 0) for item in self.current_quote.get("items", []))
        self.current_quote["total_value"] = total
        self.total_label.setText(f"${total:.2f}")

    def _on_customer_info_changed(self):
        """Handle customer information changes."""
        self.current_quote["customer_info"] = {
            "company_name": self.company_name_edit.text().strip(),
            "contact_name": self.contact_person_edit.text().strip(),
            "email": self.email_edit.text().strip(),
            "phone": self.phone_edit.text().strip(),
            "notes": self.notes_edit.toPlainText().strip()
        }
        
        # Update action button states
        self._update_items_visibility()

    def _save_draft(self):
        """Save quote as draft."""
        try:
            if not self.current_quote.get("items"):
                QMessageBox.warning(self, "No Items", "Please add at least one product to the quote.")
                return

            # Prepare quote data for saving, ensuring status is 'Draft'
            self.current_quote["status"] = "Draft"
            quote_data = {
                "id": self.current_quote.get("id"),
                "customer_info": self.current_quote.get("customer_info", {}),
                "items": self.current_quote.get("items", []),
                "total_value": self.current_quote.get("total_value", 0.0),
                "status": "Draft"
            }

            # Save to database
            saved_quote = self._save_quote_to_database(quote_data)

            if saved_quote:
                # Update the current quote state with the ID from the database
                self.current_quote["id"] = saved_quote.id
                self.current_quote["quote_number"] = saved_quote.quote_number
                self.quote_number_label.setText(f"Quote #: {saved_quote.quote_number}")

                QMessageBox.information(self, "Saved", f"Quote #{saved_quote.quote_number} saved as draft.")
                self.quote_created.emit(self.current_quote)

        except Exception as e:
            logger.error(f"Error saving quote: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Failed to save quote: {str(e)}")

    def _export_to_word(self):
        """Export quote to Word."""
        try:
            if not self.current_quote.get("items"):
                QMessageBox.warning(self, "No Items", "Please add at least one product to the quote.")
                return
            
            # Prepare quote data for export
            quote_data = self._prepare_quote_data_for_export()
            
            # Get save location
            default_filename = f"Quote_{quote_data['quote_number']}_{quote_data['customer_name'].replace(' ', '_')}.docx"
            save_dir = "data/quotes"
            os.makedirs(save_dir, exist_ok=True)
            default_path = os.path.join(save_dir, default_filename)
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Quote as Word Document",
                default_path,
                "Word Documents (*.docx)"
            )
            
            if not file_path:
                return
            
            # Generate Word document
            self._generate_word_document(quote_data, file_path)
            
            QMessageBox.information(self, "Success", f"Quote exported to:\n{file_path}")
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Failed to export quote to Word: {str(e)}")

    def _prepare_quote_data_for_export(self):
        """Prepare quote data for Word export."""
        from datetime import datetime
        
        # Get customer info
        customer_data = {
            'name': self.company_name_edit.text() or 'N/A',
            'contact_person': self.contact_person_edit.text() or 'N/A',
            'email': self.email_edit.text() or 'N/A',
            'phone': self.phone_edit.text() or 'N/A',
            'notes': self.notes_edit.toPlainText() or ''
        }
        
        # Get quote items
        items = []
        total_price = 0.0
        
        for item in self.current_quote.get("items", []):
            items.append({
                'product': item.get("product_family", "N/A"),
                'configuration': item.get("configuration", "Standard Configuration"),
                'quantity': item.get("quantity", 1),
                'unit_price': item.get("unit_price", 0.0),
                'total': item.get("total_price", 0.0)
            })
            total_price += item.get("total_price", 0.0)
        
        # Generate quote number if not exists
        quote_number = self.current_quote.get("quote_number")
        if not quote_number:
            date_str = datetime.now().strftime('%Y-%m%d')
            quote_number = f"Q-{date_str}-001"
        
        # Prepare complete quote data
        quote_data = {
            'quote_number': quote_number,
            'date': datetime.now().strftime('%Y-%m-%d'),
            'customer_name': customer_data['name'],
            'contact_person': customer_data['contact_person'],
            'subject': f"{items[0]['product']} Level Transmitter" if items else "Quote",
            'items': items,
            'total_price': total_price,
            'application_notes': self._get_application_notes(items),
            'sales_person_name': 'John Nichelosi',
            'sales_person_phone': '(713) 467-4438',
            'sales_person_email': 'John@babbitt.us',
            'company_info': {
                'name': 'Babbitt International',
                'address': 'Houston, TX',
                'contact': 'Email: sales@babbittinternational.com | Phone: (713) 467-4438',
            },
            'terms_and_conditions': (
                '1. All prices are in USD.\n'
                '2. Terms: Net 30 days W.A.C. or CC\n'
                '3. Prices are valid for 30 days.\n'
                '4. Delivery: FCA Factory, Houston, TX'
            )
        }
        
        return quote_data

    def _get_application_notes(self, items):
        """Get application notes based on products."""
        # Default LT9000 application notes from the price list
        lt9000_notes = """THE LT 9000 IS DESIGNED TO BE USED IN ELECTRICALLY CONDUCTIVE LIQUIDS THAT DO NOT LEAVE A RESIDUE ON THE PROBE. A wet electrically conductive coating will give an indication of level at the highest point that there is a continuous coating from the surface of the fluid.

For proper operation, the LT 9000 must be grounded to the fluid. In non-metallic tanks, extra grounding provisions may be necessary. It is good engineering practice to provide a separate independent high-level alarm in critical applications, rather than using a set point based on the 4-20mA output."""
        
        # Check if any LT9000 products in the quote
        for item in items:
            if 'LT9000' in item['product']:
                return lt9000_notes
        
        # Default notes for other products
        return "Please refer to product manual for detailed application notes and installation instructions."

    def _generate_word_document(self, quote_data, output_path):
        """Generate a Word document from quote data."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Header with company name
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = ""
            run = header_para.add_run("BABBITT\nINTERNATIONAL")
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 139)
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add subtitle
            subtitle_para = header.add_paragraph()
            subtitle_run = subtitle_para.add_run("Level Controls & Systems")
            subtitle_run.font.name = 'Arial'
            subtitle_run.font.size = Pt(10)
            subtitle_run.font.italic = True
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Main document content
            doc.add_paragraph(f"DATE: {quote_data['date']}")
            doc.add_paragraph(f"CUSTOMER: {quote_data['customer_name']}")
            doc.add_paragraph(f"ATTN: {quote_data['contact_person']}")
            
            # Quote header
            quote_header = doc.add_paragraph()
            quote_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            quote_run = quote_header.add_run(f"Quote # {quote_data['quote_number']}")
            quote_run.bold = True
            
            # Subject line
            subject = doc.add_paragraph()
            subject_run = subject.add_run(f"Subject: {quote_data['subject']}")
            subject_run.bold = True
            
            # Introduction text
            intro = doc.add_paragraph(
                "We are pleased to quote on the following equipment for your upcoming applications:"
            )
            
            # Add spacing
            doc.add_paragraph()
            
            # Quote items table
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'QTY'
            header_cells[1].text = 'PRODUCT'
            header_cells[2].text = 'CONFIGURATION'
            header_cells[3].text = 'UNIT PRICE'
            header_cells[4].text = 'TOTAL'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for item in quote_data['items']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['quantity'])
                row_cells[1].text = item['product']
                row_cells[2].text = item['configuration']
                row_cells[3].text = f"${item['unit_price']:,.2f}"
                row_cells[4].text = f"${item['total']:,.2f}"
            
            # Add total row
            total_row = table.add_row()
            total_cells = total_row.cells
            total_cells[0].text = ""
            total_cells[1].text = ""
            total_cells[2].text = ""
            total_cells[3].text = "TOTAL:"
            total_cells[4].text = f"${quote_data['total_price']:,.2f}"
            
            # Make total row bold
            for cell in total_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Delivery and terms
            delivery = doc.add_paragraph("Delivery:")
            delivery.add_run("\nTerms: Net 30 days W.A.C. or CC")
            
            validity = doc.add_paragraph(
                "FCA: Factory, Houston, TX\nQuotation valid for 30 days."
            )
            
            # Application notes section
            doc.add_paragraph()
            app_notes_header = doc.add_paragraph("APPLICATION NOTES")
            app_notes_header.runs[0].bold = True
            app_notes_header.runs[0].underline = True
            
            # Application notes
            doc.add_paragraph(quote_data['application_notes'])
            
            # Add spacing before footer
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Footer text
            footer_text = doc.add_paragraph(
                "Please contact me directly if you have any questions or require more information."
            )
            
            doc.add_paragraph("Thank you,")
            doc.add_paragraph()
            doc.add_paragraph(quote_data['sales_person_name'])
            doc.add_paragraph(quote_data['sales_person_phone'])
            doc.add_paragraph(quote_data['sales_person_email'])
            doc.add_paragraph("www.babbittinternational.com")
            
            # Save document
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}", exc_info=True)
            raise Exception(f"Failed to generate Word document: {str(e)}")

    def _save_quote_to_database(self, quote_data: Dict) -> Optional['Quote']:
        """Save quote to database and return the full Quote object."""
        try:
            with SessionLocal() as db:
                # Prepare customer data
                customer_info = quote_data.get("customer_info", {})
                
                # Prepare product data
                products_data = []
                for item in quote_data.get("items", []):
                    products_data.append({
                        "product_id": item.get("product_id"),
                        "quantity": item.get("quantity", 1),
                        "base_price": item.get("unit_price", 0),
                        "part_number": item.get("model_number", "N/A"),
                        "options": item.get("options", [])
                    })

                quote_details = {
                    "notes": customer_info.get("notes")
                }

                # If quote has an ID, it's an update. Otherwise, it's a new quote.
                if self.current_quote_id is not None:
                    quote = self.quote_service.update_quote_with_items(
                        db,
                        quote_id=self.current_quote_id,
                        customer_data=customer_info,
                        products_data=products_data,
                        quote_details=quote_details,
                    )
                else:
                    quote = self.quote_service.create_quote_with_items(
                        db,
                        customer_data=customer_info,
                        products_data=products_data,
                        quote_details=quote_details,
                    )
                
                # After saving, get the updated quote ID and number
                db.refresh(quote)
                self.current_quote["id"] = quote.id
                self.current_quote["quote_number"] = quote.quote_number
                
                # This is the critical fix: update the tracked quote ID
                self.current_quote_id = quote.id

                logger.info(f"Successfully saved quote {quote.quote_number} with ID {quote.id}")
                return quote
                
        except Exception as e:
            logger.error(f"Error saving to database: {e}", exc_info=True)
            QMessageBox.critical(self, "Database Error", f"Could not save quote: {e}")
            return None

    def new_quote(self):
        """Start a new quote (clear current state)."""
        reply = QMessageBox.question(
            self,
            "New Quote",
            "Are you sure you want to start a new quote? Any unsaved changes will be lost.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self._reset_quote()

    def _reset_quote(self):
        """Reset the quote to initial state."""
        self.current_quote = {
            "items": [],
            "customer_info": {},
            "total_value": 0.0,
            "quote_number": None,
            "status": "Draft"
        }
        
        # Clear UI
        self.items_table.setRowCount(0)
        self.company_name_edit.clear()
        self.contact_person_edit.clear()
        self.email_edit.clear()
        self.phone_edit.clear()
        self.notes_edit.clear()
        
        # Reset labels
        self.quote_number_label.setText("Quote not yet saved")
        self.status_label.setText("DRAFT")
        self.status_label.setProperty("status", "Draft")
        self.total_label.setText("$0.00")
        
        self._update_items_visibility()

        # Force re-application of stylesheet to fix styling glitches
        self.style().unpolish(self)
        self.style().polish(self)

    def load_quote(self, quote_data: Dict):
        """Loads an existing quote into the editor."""
        self._reset_quote()
        
        self.current_quote_id = quote_data.get("id")
        self.current_quote = {
            "items": quote_data.get("products", []),
            "customer_info": quote_data.get("customer", {}),
            "total_value": quote_data.get("total", 0.0),
            "quote_number": quote_data.get("quote_number"),
            "status": quote_data.get("status", "Draft"),
            "notes": quote_data.get("notes")
        }
        
        self._update_ui_from_quote()
        logger.info(f"Loaded quote {self.current_quote_id} into creator.")

    def _update_ui_from_quote(self):
        """Populates all UI fields from the self.current_quote data."""
        # Populate customer info
        customer_info = self.current_quote.get("customer_info", {})
        self.company_name_edit.setText(customer_info.get("company_name", customer_info.get("company", "")))
        self.contact_person_edit.setText(customer_info.get("contact_name", customer_info.get("name", "")))
        self.email_edit.setText(customer_info.get("email", ""))
        self.phone_edit.setText(customer_info.get("phone", ""))
        self.notes_edit.setPlainText(self.current_quote.get("notes", ""))
        
        # Update quote header labels
        quote_number = self.current_quote.get("quote_number")
        if quote_number:
            self.quote_number_label.setText(f"Quote #: {quote_number}")
        else:
            self.quote_number_label.setText("Quote not yet saved")

        status = self.current_quote.get("status", "Draft")
        self.status_label.setText(status.upper())
        self.status_label.setProperty("status", status)
        self.status_label.style().unpolish(self.status_label)
        self.status_label.style().polish(self.status_label)
        
        # Update items table and totals
        self._update_items_table()
        self._update_quote_totals()
        self._update_items_visibility()

    def clear_if_quote_matches(self, deleted_quote_id: int):
        """Clears the form if the deleted quote ID matches the current one."""
        if self.current_quote_id and self.current_quote_id == deleted_quote_id:
            logger.info(f"Clearing quote creator because quote {deleted_quote_id} was deleted.")
            self._reset_quote()

    def closeEvent(self, event):
        """Handle the window closing."""
        if hasattr(self, 'db') and self.db:
            self.db.close()
        event.accept()

    def clear_form(self, deleted_quote_id: int):
        """
        Clears the form if the deleted quote matches the one currently loaded.
        """
        if self.current_quote_id == deleted_quote_id:
            logger.info(f"Quote {deleted_quote_id} was deleted while being edited. Clearing form.")
            self.new_quote()
            QMessageBox.information(
                self, 
                "Quote Deleted", 
                "The quote you were editing has been deleted. The form has been reset."
            )