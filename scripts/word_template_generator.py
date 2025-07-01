"""
Script to generate the Word template for quotes.
Run this once to create the template file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def add_table_borders(table):
    """Add borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Create border element
            tcBorders = OxmlElement('w:tcBorders')
            
            # Add borders to all sides
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)


def create_quote_template():
    """Create a Word template for quotes."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header with company name
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "" # Clear existing text
    run = header_para.add_run("BABBITT\nINTERNATIONAL")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add subtitle
    subtitle_para = header.add_paragraph()
    subtitle_run = subtitle_para.add_run("Level Controls & Systems")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Main document content
    # Date and customer info section
    doc.add_paragraph("DATE: {{date}}")
    doc.add_paragraph("CUSTOMER: {{customer_name}}")
    doc.add_paragraph("ATTN: {{contact_person}}")
    
    # Quote header
    quote_header = doc.add_paragraph()
    quote_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    quote_run = quote_header.add_run("Quote # {{quote_number}}")
    quote_run.bold = True
    
    # Subject line
    subject = doc.add_paragraph()
    subject_run = subject.add_run("Subject: {{subject}}")
    subject_run.bold = True
    
    # Introduction text
    intro = doc.add_paragraph(
        "We are pleased to quote on the following equipment for your upcoming applications:"
    )
    
    # Add spacing
    doc.add_paragraph()
    
    # Quote items table
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'QTY'
    header_cells[1].text = 'PRODUCT'
    header_cells[2].text = 'CONFIGURATION'
    header_cells[3].text = 'UNIT PRICE'
    header_cells[4].text = 'TOTAL'
    header_cells[5].text = 'EACH'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Sample data row (will be replaced by actual data)
    data_cells = table.rows[1].cells
    data_cells[0].text = '{{quantity}}'
    data_cells[1].text = '{{product_code}}'
    data_cells[2].text = '{{configuration}}'
    data_cells[3].text = '$'
    data_cells[4].text = '{{unit_price}}'
    data_cells[5].text = 'EACH'
    
    # Add borders to table
    add_table_borders(table)
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Delivery and terms
    delivery = doc.add_paragraph("Delivery:")
    delivery.add_run("\nTerms: Net 30 days W.A.C. or CC")
    
    validity = doc.add_paragraph(
        "FCA: Factory, Houston, TX\nQuotation valid for 30 days."
    )
    
    # Application notes section
    doc.add_paragraph()
    app_notes_header = doc.add_paragraph("APPLICATION NOTES")
    app_notes_header.runs[0].bold = True
    app_notes_header.runs[0].underline = True
    
    # Placeholder for application notes
    doc.add_paragraph("{{application_notes}}")
    
    # Add spacing before footer
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer text
    footer_text = doc.add_paragraph(
        "Please contact me directly if you have any questions or require more information."
    )
    
    doc.add_paragraph("Thank you,")
    doc.add_paragraph()
    doc.add_paragraph("{{sales_person_name}}")
    doc.add_paragraph("{{sales_person_phone}}")
    doc.add_paragraph("{{sales_person_email}}")
    doc.add_paragraph("www.babbittinternational.com")
    
    # Save template
    template_dir = "data/templates"
    os.makedirs(template_dir, exist_ok=True)
    
    template_path = os.path.join(template_dir, "quote_template.docx")
    doc.save(template_path)
    
    print(f"Quote template created at: {template_path}")
    return template_path


def generate_quote_from_template(template_path, output_path, context):
    """
    Generate a quote from a Word template by replacing placeholders.
    Enhanced version that handles complex template data structures.
    """
    try:
        doc = Document(template_path)
        
        # Debug: Print context keys to see what's available
        print(f"🔍 Available context keys: {list(context.keys())}")
        
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            if "{{" in original_text and "}}" in original_text:
                print(f"🔍 Found paragraph with placeholders: {original_text}")
                
                # Replace all placeholders in this paragraph
                new_text = original_text
                for key, value in context.items():
                    # Skip complex data structures like lists
                    if isinstance(value, (list, dict)):
                        continue
                    
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in new_text:
                        print(f"  🔄 Replacing {placeholder} with '{value}'")
                        new_text = new_text.replace(placeholder, str(value))
                
                # Update the paragraph text if it changed
                if new_text != original_text:
                    para.text = new_text
                    print(f"  ✅ Updated paragraph: {new_text}")
        
        # Replace placeholders in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        original_text = para.text
                        if "{{" in original_text and "}}" in original_text:
                            print(f"🔍 Found table cell with placeholders: {original_text}")
                            
                            # Replace all placeholders in this paragraph
                            new_text = original_text
                            for key, value in context.items():
                                # Skip complex data structures like lists
                                if isinstance(value, (list, dict)):
                                    continue
                                
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in new_text:
                                    print(f"  🔄 Replacing {placeholder} with '{value}'")
                                    new_text = new_text.replace(placeholder, str(value))
                            
                            # Update the paragraph text if it changed
                            if new_text != original_text:
                                para.text = new_text
                                print(f"  ✅ Updated table cell: {new_text}")

        # Replace placeholders in headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                original_text = para.text
                if "{{" in original_text and "}}" in original_text:
                    print(f"🔍 Found header with placeholders: {original_text}")
                    
                    # Replace all placeholders in this paragraph
                    new_text = original_text
                    for key, value in context.items():
                        # Skip complex data structures like lists
                        if isinstance(value, (list, dict)):
                            continue
                        
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in new_text:
                            print(f"  🔄 Replacing {placeholder} with '{value}'")
                            new_text = new_text.replace(placeholder, str(value))
                    
                    # Update the paragraph text if it changed
                    if new_text != original_text:
                        para.text = new_text
                        print(f"  ✅ Updated header: {new_text}")

        # Replace placeholders in footers
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                original_text = para.text
                if "{{" in original_text and "}}" in original_text:
                    print(f"🔍 Found footer with placeholders: {original_text}")
                    
                    # Replace all placeholders in this paragraph
                    new_text = original_text
                    for key, value in context.items():
                        # Skip complex data structures like lists
                        if isinstance(value, (list, dict)):
                            continue
                        
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in new_text:
                            print(f"  🔄 Replacing {placeholder} with '{value}'")
                            new_text = new_text.replace(placeholder, str(value))
                    
                    # Update the paragraph text if it changed
                    if new_text != original_text:
                        para.text = new_text
                        print(f"  ✅ Updated footer: {new_text}")

        doc.save(output_path)
        print(f"✅ Quote generated at: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error generating quote: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    # 1. Ensure the template exists
    template_file = "data/templates/quote_template.docx"
    if not os.path.exists(template_file):
        print("Template not found, creating a new one...")
    create_quote_template()

    # 2. Define sample context data
    sample_context = {
        "date": "2025-06-24",
        "customer_name": "ACME Corp",
        "contact_person": "Wile E. Coyote",
        "quote_number": "Q-2025-0624-001",
        "subject": "LT9000 Level Transmitter",
        "quantity": 1,
        "product_code": "LT9000-XXXX-H-XX",
        "configuration": '• Continuous Level Transmitter\n'
                         '• Supply Voltage: 4 to 20mA\n'
                         '• Process Connection: 1" NPT; 316SS (1500 PSI Max)\n'
                         '• Insulator: Teflon, 4" Long (350 F)\n'
                         '• Probe: ½" Diameter HALAR x XX" (Including Insulator)\n'
                         '• Housing: Cast Aluminum, NEMA 7, D; NEMA 9, E, F, & G\n'
                         '• 2-Year Warranty',
        "unit_price": "1,234.56",
        "application_notes": "THE LT 9000 IS DESIGNED TO BE USED IN ELECTRICALLY CONDUCTIVE LIQUIDS THAT DO NOT LEAVE A RESIDUE ON THE PROBE. A wet electrically conductive coating will give an indication of level at the highest point that there is a continuous coating from the surface of the fluid. The LT 9000 will give a varying output, if the conductivity of the material changes. For proper operation, the LT 9000 must be grounded to the fluid. In non-metallic tanks, extra grounding provisions may be necessary. It is good engineering practice to provide a separate independent high-level alarm in critical applications, rather than using a set point based on the 4-20mA output.",
        "sales_person_name": "John Nicholosi",
        "sales_person_phone": "(713) 467-4438",
        "sales_person_email": "john@babbitt.us",
    }

    # 3. Generate the quote
    output_file = f"Quote_{sample_context['quote_number']}_{sample_context['customer_name']}.docx"
    generate_quote_from_template(template_file, output_file, sample_context)
